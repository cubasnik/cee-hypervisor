import json
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime

def create_styled_documentation():
    # Загрузка спецификации
    print("📥 Загрузка спецификации...")
    response = requests.get('http://172.30.226.25:8001/openapi.json')
    spec = response.json()
    
    # Создание документа
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Титульная страница
    title = doc.add_heading(f"{spec['info']['title']}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_heading('Документация API', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph(f"Версия: {spec['info']['version']}")
    doc.add_paragraph(f"Дата создания: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}")
    
    if 'description' in spec['info']:
        doc.add_paragraph()
        doc.add_heading('Описание', level=2)
        doc.add_paragraph(spec['info']['description'])
    
    doc.add_page_break()
    
    # Содержание
    doc.add_heading('Содержание', level=1)
    toc = doc.add_paragraph()
    toc.add_run('1. Серверы').bold = True
    toc.add_run('\n2. Эндпоинты').bold = True
    toc.add_run('\n3. Модели данных').bold = True
    
    doc.add_page_break()
    
    # Серверы
    if 'servers' in spec:
        doc.add_heading('1. Серверы', level=1)
        for server in spec['servers']:
            p = doc.add_paragraph(style='List Number')
            p.add_run(f"URL: {server.get('url', '')}").bold = True
            if 'description' in server:
                doc.add_paragraph(server['description'], style='Intense Quote')
    
    doc.add_page_break()
    
    # Эндпоинты
    doc.add_heading('2. Эндпоинты', level=1)
    
    # Группировка по тегам
    endpoints_by_tag = {}
    for path, methods in spec.get('paths', {}).items():
        for method, details in methods.items():
            tags = details.get('tags', ['General'])
            for tag in tags:
                if tag not in endpoints_by_tag:
                    endpoints_by_tag[tag] = []
                endpoints_by_tag[tag].append({
                    'path': path,
                    'method': method,
                    'details': details
                })
    
    for tag, endpoints in endpoints_by_tag.items():
        doc.add_heading(f'2.{list(endpoints_by_tag.keys()).index(tag) + 1} {tag}', level=2)
        
        for endpoint in endpoints:
            # Заголовок эндпоинта
            heading = doc.add_heading(level=3)
            run = heading.add_run(f"{endpoint['method'].upper()} {endpoint['path']}")
            run.font.color.rgb = RGBColor(0, 102, 204)
            
            # Сводка
            if 'summary' in endpoint['details']:
                doc.add_paragraph(endpoint['details']['summary'], style='Intense Quote')
            
            # Параметры
            if 'parameters' in endpoint['details'] and endpoint['details']['parameters']:
                doc.add_heading('Параметры', level=4)
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Light Grid Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Имя'
                hdr_cells[1].text = 'Тип'
                hdr_cells[2].text = 'Обязательный'
                hdr_cells[3].text = 'Описание'
                
                for param in endpoint['details']['parameters']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = param.get('name', '')
                    row_cells[1].text = param.get('schema', {}).get('type', 'string')
                    row_cells[2].text = '✓' if param.get('required', False) else '✗'
                    row_cells[3].text = param.get('description', '')
            
            # Request Body
            if 'requestBody' in endpoint['details']:
                doc.add_heading('Тело запроса', level=4)
                content = endpoint['details']['requestBody'].get('content', {})
                for content_type, content_details in content.items():
                    doc.add_paragraph(f"Content-Type: {content_type}", style='List Bullet')
                    if 'schema' in content_details:
                        schema = content_details['schema']
                        doc.add_paragraph(f"Схема: {json.dumps(schema, indent=2, ensure_ascii=False)}", style='Code')
            
            # Ответы
            if 'responses' in endpoint['details']:
                doc.add_heading('Ответы', level=4)
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light Grid Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Код'
                hdr_cells[1].text = 'Описание'
                
                for status, response in endpoint['details']['responses'].items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = status
                    row_cells[1].text = response.get('description', '')
            
            doc.add_paragraph()  # Пустая строка
    
    doc.add_page_break()
    
    # Модели данных
    doc.add_heading('3. Модели данных', level=1)
    if 'components' in spec and 'schemas' in spec['components']:
        for schema_name, schema in spec['components']['schemas'].items():
            doc.add_heading(f"3.{list(spec['components']['schemas'].keys()).index(schema_name) + 1} {schema_name}", level=2)
            
            if 'description' in schema:
                doc.add_paragraph(schema['description'])
            
            if 'properties' in schema:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Light Grid Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Поле'
                hdr_cells[1].text = 'Тип'
                hdr_cells[2].text = 'Обязательное'
                hdr_cells[3].text = 'Описание'
                
                required = schema.get('required', [])
                for prop_name, prop in schema['properties'].items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = prop_name
                    row_cells[1].text = prop.get('type', 'string')
                    row_cells[2].text = '✓' if prop_name in required else '✗'
                    row_cells[3].text = prop.get('description', '')
    
    # Сохранение
    filename = f"cee-hypervisor-api-docs_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
    doc.save(filename)
    print(f"✅ Документация сохранена в {filename}")

if __name__ == "__main__":
    create_styled_documentation()
